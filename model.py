import math
from geopy.distance import geodesic
from fastkml import kml
from shapely.geometry import shape
import csv
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import numpy as np
from sklearn.linear_model import LinearRegression
import ee
import geemap
import matplotlib.pyplot as plt
from sklearn.linear_model import LinearRegression
import numpy as np

# Function to calculate area of a geodesic polygon
def calculate_polygon_area(vertices):
    R = 6378137  # Radius of the Earth in meters
    vertices = [(math.radians(lat), math.radians(lon)) for lat, lon in vertices]
    area = 0
    for i in range(len(vertices)):
        lat1, lon1 = vertices[i]
        lat2, lon2 = vertices[(i + 1) % len(vertices)]
        area += (lon2 - lon1) * (2 + math.sin(lat1) + math.sin(lat2))
    return abs(area) / 2.0 * (R**2)

# Function to parse .csv file and extract lake coordinates with full details
def parse_csv_lakes(file_path):
    lakes = []
    
    with open(file_path, "r", newline='', encoding='utf-8') as csv_file:
        reader = csv.DictReader(csv_file)  # Read CSV into a dictionary
        for row in reader:
            # Extract the required details
            name_of_th = row.get('Name_of_Th')
            unique_id = row.get('UniqueID')
            valley = row.get('Valley')
            area = row.get('Area')
            latitude = row.get('Latitude')
            longitude = row.get('Longitude')
            ward_number = row.get('Ward_Numbe')
            ward_name = row.get('Ward_Name')
            ward_counc = row.get('Ward_Couns')
            ward_office = row.get('Ward_offic')
            aream3 = row.get('Aream3')
            
            # Only add if coordinates are valid
            if latitude and longitude:
                try:
                    lat = float(latitude)
                    lon = float(longitude)
                    lake = {
                        'name_of_th': name_of_th,
                        'unique_id': unique_id,
                        'valley': valley,
                        'area': area,
                        'latitude': lat,
                        'longitude': lon,
                        'ward_number': ward_number,
                        'ward_name': ward_name,
                        'ward_counc': ward_counc,
                        'ward_office': ward_office,
                        'aream3': aream3
                    }
                    lakes.append(lake)
                except ValueError:
                    continue  # Skip rows where latitude or longitude are not valid floats
                    
    return lakes

# Function to check if plot is within buffer zone and return closest lake details
def is_within_buffer(plot_vertices, lake_coords, lakes, min_buffer=30, max_buffer=75):
    plot_center = (
        sum([lat for lat, _ in plot_vertices]) / len(plot_vertices),
        sum([lon for _, lon in plot_vertices]) / len(plot_vertices),
    )
    
    closest_lake = None
    closest_distance = float('inf')
    
    # Loop through the lakes to find the closest one
    for i, lake_coord in enumerate(lake_coords):
        distance = geodesic(plot_center, lake_coord).meters
        if distance < closest_distance:
            closest_distance = distance
            closest_lake = lakes[i]  # Get the corresponding lake details

        if min_buffer <= distance <= max_buffer:
            return True, distance, closest_lake
    
    return False, closest_distance, closest_lake

# Function to display details of the closest lake
def get_closest_lake_details(closest_lake):
    if closest_lake:
        print(f"Closest Lake Details:")
        print(f"Name: {closest_lake['name_of_th']}")
        print(f"Unique ID: {closest_lake['unique_id']}")
        print(f"Area: {closest_lake['area']}")
        print(f"Latitude: {closest_lake['latitude']}")
        print(f"Longitude: {closest_lake['longitude']}")
        print(f"Ward Number: {closest_lake['ward_number']}")
        print(f"Ward Name: {closest_lake['ward_name']}")
        print(f"Ward Councillor: {closest_lake['ward_counc']}")
        print(f"Ward Office: {closest_lake['ward_office']}")
        print(f"Area (m^3): {closest_lake['aream3']}")
    else:
        print("No closest lake found.")

# Authenticate and initialize the Earth Engine API
ee.Authenticate()
ee.Initialize()

# Function to compute NDVI
def calculate_ndvi(image):
    ndvi = image.normalizedDifference(['B8', 'B4']).rename('NDVI')  # Updated band names
    return image.addBands(ndvi)

# Function to calculate current GCI
def calculate_current_gci(coordinates):
    # Automatically close the polygon by appending the first point to the end
    if coordinates[0] != coordinates[-1]:
        coordinates.append(coordinates[0])  # Close the polygon

    # Define the area of interest (AOI) as a polygon
    aoi = ee.Geometry.Polygon([coordinates])
    print("AOI Geometry:", aoi.getInfo())


    # Fetch Sentinel-2 imagery for the recent time range
    image_collection = ee.ImageCollection('COPERNICUS/S2_SR_HARMONIZED') \
        .filterBounds(aoi) \
        .filterDate('2025-01-20', '2025-01-25') \
        .map(calculate_ndvi) \
        .select('NDVI')

    # Calculate the mean NDVI within the AOI
    mean_ndvi = image_collection.mean().reduceRegion(
        reducer=ee.Reducer.mean(),
        geometry=aoi,
        scale=30,
        maxPixels=1e8
    )

    # Extract the mean NDVI value (Green Cover Index)
    gci = mean_ndvi.get('NDVI').getInfo()
    return gci

# Function to calculate yearly GCI (Green Cover Index)
def calculate_yearly_gci(coordinates):
    # Automatically close the polygon by appending the first point to the end
    if coordinates[0] != coordinates[-1]:
        coordinates.append(coordinates[0])  # Close the polygon

    # Define the area of interest (AOI) as a polygon
    aoi = ee.Geometry.Polygon([coordinates])
    
    # Define the date range for yearly calculations (April 1, 2017 to January 1, 2025)
    start_date = '2017-04-01'
    end_date = '2025-01-01'
    
    # Create an array of years to iterate over
    years = np.arange(2017, 2025)

    # List to hold GCI values for each year
    gci_values = []

    # Loop through each year to calculate GCI
    for year in years:
        # Define the start and end date for the current year
        year_start = f'{year}-04-01'
        year_end = f'{year+1}-04-01'
        
        # Fetch Sentinel-2 imagery for the current year
        image_collection = ee.ImageCollection('COPERNICUS/S2_SR_HARMONIZED') \
            .filterBounds(aoi) \
            .filterDate(year_start, year_end) \
            .map(calculate_ndvi) \
            .select('NDVI')
        
        # Calculate the mean NDVI within the AOI for this year
        mean_ndvi = image_collection.mean().reduceRegion(
            reducer=ee.Reducer.mean(),
            geometry=aoi,
            scale=30,
            maxPixels=1e8
        )

        # Extract the mean NDVI value (GCI)
        gci = mean_ndvi.get('NDVI').getInfo()
        
        # Append the GCI value to the list
        gci_values.append(gci)
    
    return years, gci_values

# Function to read plot points from a CSV file
def get_plot_points_from_csv(file_path):
    plot_points = []
    plot_points_gci = []  # Flipped coordinates for GCI

    try:
        with open(file_path, "r", newline='', encoding='utf-8') as csv_file:
            reader = csv.DictReader(csv_file)
            for row in reader:
                latitude = row.get('Latitude')
                longitude = row.get('Longitude')

                if latitude and longitude:
                    try:
                        lat = float(latitude)
                        lon = float(longitude)
                        plot_points.append((lat, lon))
                        plot_points_gci.append((lon, lat))  # Flip coordinates for GCI calculation
                    except ValueError:
                        print(f"Invalid coordinates: {latitude}, {longitude}. Skipping row.")
    except FileNotFoundError:
        print(f"Error: File not found - {file_path}")
    except Exception as e:
        print(f"Error reading CSV file: {e}")

    # print(plot_points, plot_points_gci)
    return plot_points, plot_points_gci

from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import numpy as np

def save_report_to_word(file_path, years, gci_values, future_years, predicted_gci, plot_points, area, within_buffer, distance, closest_lake, output_file="report.docx"):
    # Create a new Word document
    doc = Document()

    # Add a title
    doc.add_heading("Green Cover Index (GCI) Analysis Report", level=1)

    # Add a section for the introduction
    doc.add_heading("Introduction", level=2)
    doc.add_paragraph(
        "This report provides an analysis of the Green Cover Index (GCI) over the years, predicts its trends for the next decade, and includes details about the plot's area and proximity to lakes."
    )

    # Add a section for plot details
    doc.add_heading("Plot Analysis", level=2)

    # Plot points
    doc.add_paragraph("The following are the coordinates of the plot points:")
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Latitude"
    hdr_cells[1].text = "Longitude"
    for lat, lon in plot_points:
        row_cells = table.add_row().cells
        row_cells[0].text = f"{lat:.6f}"
        row_cells[1].text = f"{lon:.6f}"

    # Calculated area
    doc.add_paragraph(f"The calculated area of the plot is {area:.2f} square meters.")

    # Buffer zone status
    if within_buffer:
        doc.add_paragraph(f"The plot is within the legal buffer zone and is {distance:.2f} meters away from the nearest lake.")
    else:
        doc.add_paragraph(f"The plot is outside the legal buffer zone. The closest lake details are as follows:")
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Attribute"
        hdr_cells[1].text = "Value"
        for key, value in closest_lake.items():
            row_cells = table.add_row().cells
            row_cells[0].text = str(key)
            row_cells[1].text = str(value)

    # Add a section for historical GCI
    doc.add_heading("Historical GCI Data", level=2)
    doc.add_paragraph("The following table shows the historical GCI values:")

    # Add a table for historical data
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Year"
    hdr_cells[1].text = "GCI Value"
    for year, gci in zip(years, gci_values):
        row_cells = table.add_row().cells
        row_cells[0].text = str(year)
        row_cells[1].text = f"{gci:.2f}"

    # Add a section for the historical GCI plot
    doc.add_heading("Historical GCI Plot", level=2)
    plt.figure(figsize=(10, 6))
    plt.plot(years, gci_values, marker='o', color='b', linestyle='-', markersize=6, label='Historical GCI')
    plt.title('Yearly Green Cover Index (GCI)')
    plt.xlabel('Year')
    plt.ylabel('GCI')
    plt.grid(True)
    plt.legend()
    plt.tight_layout()
    plt.savefig("historical_gci.png")
    plt.close()

    doc.add_paragraph("The following plot shows the historical GCI trends:")
    doc.add_picture("historical_gci.png", width=Inches(5))

    # Add a section for GCI predictions
    doc.add_heading("GCI Predictions (2025-2035)", level=2)
    doc.add_paragraph("Using a linear regression model, the GCI values for the next decade have been predicted:")

    # Add a table for predicted data
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Year"
    hdr_cells[1].text = "Predicted GCI Value"
    for year, gci in zip(future_years.flatten(), predicted_gci):
        row_cells = table.add_row().cells
        row_cells[0].text = str(year)
        row_cells[1].text = f"{gci:.2f}"

    # Add a section for the prediction plot
    doc.add_heading("GCI Predictions Plot", level=2)
    plt.figure(figsize=(10, 6))
    plt.plot(years, gci_values, marker='o', color='b', linestyle='-', markersize=6, label='Historical GCI')
    plt.plot(future_years, predicted_gci, marker='x', color='r', linestyle='--', markersize=6, label='Predicted GCI')
    plt.title('Yearly Green Cover Index (GCI) and Predictions (2025-2035)')
    plt.xlabel('Year')
    plt.ylabel('GCI')
    plt.grid(True)
    plt.xticks(np.arange(min(years), max(future_years.flatten()) + 1, 1))
    plt.legend()
    plt.tight_layout()
    plt.savefig("predicted_gci.png")
    plt.close()

    doc.add_paragraph("The following plot shows the historical GCI values and the predictions for 2025-2035:")
    doc.add_picture("predicted_gci.png", width=Inches(5))

    # Save the document
    doc.save(output_file)
    print(f"Report saved as {output_file}")


# Example usage of `get_plot_points_from_csv` and related functionality
def main(file_path):
    # Read plot points and flipped coordinates for GCI
    plot_points, plot_points_gci = get_plot_points_from_csv(file_path)
    
    # If no points were found, exit
    if not plot_points:
        print("No valid plot points found in the CSV file.")
        return

    # Example: Calculate area of the polygon formed by the plot points
    area = calculate_polygon_area(plot_points)
    print(f"Calculated area of the plot: {area:.2f} square meters")

    csv_file_path = "output_lake.csv"  # Ensure this file exists
    try:
        lakes = parse_csv_lakes(csv_file_path)
        lake_coordinates = [(lake['latitude'], lake['longitude']) for lake in lakes]
        print(f"Loaded {len(lake_coordinates)} lake boundary points.")

        within_buffer, distance, closest_lake = is_within_buffer(plot_points, lake_coordinates, lakes)
        if within_buffer:
            print(f"The plot is {distance:.2f} meters away from the lake and within the buffer zone.")
        else:
            print(f"The plot is outside the legal buffer zone.")
            get_closest_lake_details(closest_lake)
    
    except Exception as e:
        print(f"Error parsing lakes from CSV: {e}")


    # Example: Compute current GCI using the flipped coordinates
    try:
        current_gci = calculate_current_gci(plot_points_gci)
        print(f"Current Green Cover Index (GCI): {current_gci}")
    except Exception as e:
        print(f"Error calculating GCI: {e}")

    # Example: Compute yearly GCI trends
    try:
        years, gci_values = calculate_yearly_gci(plot_points_gci)
        print("Yearly GCI values:")
        for year, gci in zip(years, gci_values):
            print(f"Year {year}: GCI = {gci}")
    except Exception as e:
        print(f"Error calculating yearly GCI: {e}")

    # Example: Plot yearly GCI trends
    # try:
    #     plt.figure(figsize=(10, 6))
    #     plt.plot(years, gci_values, marker='o', label='Yearly GCI')
    #     plt.title("Yearly Green Cover Index Trends")
    #     plt.xlabel("Year")
    #     plt.ylabel("Green Cover Index (GCI)")
    #     plt.grid(True)
    #     plt.legend()
    #     plt.show()
    # except Exception as e:
    #     print(f"Error plotting GCI trends: {e}")

    
# Plot historical and predicted GCI
    plt.figure(figsize=(10, 6))
    plt.plot(years, gci_values, marker='o', label='Historical GCI')
    plt.title('Yearly Green Cover Index (GCI)')
    plt.xlabel('Year')
    plt.ylabel('GCI')
    plt.grid(True)
    plt.legend()
    plt.show()

    # Convert years and GCI values to numpy arrays for regression
    years_array = np.array(years).reshape(-1, 1)  # Reshaping for regression
    gci_values_array = np.array(gci_values)

    # Fit a linear regression model to the historical data
    model = LinearRegression()
    model.fit(years_array, gci_values_array)

    # Predict GCI for the next 10 years (2025-2035)
    future_years = np.arange(2025, 2036).reshape(-1, 1)
    predicted_gci = model.predict(future_years)

    # Plot the historical GCI values and the predictions
    plt.figure(figsize=(10, 6))

    # Plot historical GCI values
    plt.plot(years, gci_values, marker='o', color='b', linestyle='-', markersize=6, label='Historical GCI')

    # Plot predicted GCI values
    plt.plot(future_years, predicted_gci, marker='x', color='r', linestyle='--', markersize=6, label='Predicted GCI')

    plt.title('Yearly Green Cover Index (GCI) and Predictions (2025-2035)')
    plt.xlabel('Year')
    plt.ylabel('Green Cover Index (GCI)')
    plt.grid(True)
    plt.xticks(np.arange(2017, 2036, 1))
    plt.legend()
    plt.tight_layout()
    plt.show()

    # Generate and save the report
    save_report_to_word(file_path, years, gci_values, future_years, predicted_gci, plot_points, area, within_buffer, distance, closest_lake)

# File path to the CSV containing plot points
file_path = "coordinates.csv"
# Allow the script to run standalone if needed
if __name__ == "__main__":
    main(file_path)